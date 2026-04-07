import hashlib
import os
import re
import tempfile
import traceback
from io import BytesIO
from logging import getLogger
from math import log
from zipfile import ZipFile

from celery import shared_task
from django.core.files.images import ImageFile
from django.db import IntegrityError, transaction
from django.db.models import Q
from docx import Document

from core.models import Question, QuestionComment, QuestionOption, QuestionImage
from courses.models import Course, Enrolment, Unit, UnitSubtopic
from .docx.parser import parse_questions_from_docx
from .docx.formats import docx_table_format_a
from .csv.parser import parse_questions_from_csv
from core.tasks.docx.parser1AA3Q import parse
from core.tasks.docx.parser1AA3exp import parse_explanation_updates
from core.tasks.upload_result_util import finish_upload_result, get_upload_result

logger = getLogger(__name__)
decimal_pattern = re.compile(r"\d?\.\d{0,5}")

PROGRESS_UPDATE_INTERVAL = 0.1


class QuestionImportError(Exception):
    """Raised when uploaded question data cannot be interpreted (any source format)."""


def _insert_question_with_logging(
    question_data: dict, source_label: str, insert_callable
) -> bool:
    serial = question_data.get("serial_number")
    try:
        insert_callable()
        logger.info(
            "Successfully inserted question serial=%s from %s",
            serial,
            source_label,
        )
        return True
    except Exception as e:
        summary = None
        if isinstance(e, IntegrityError):
            logger.error(
                "Insertion failed serial=%s (%s): integrity: %s",
                serial,
                source_label,
                e,
            )
        elif isinstance(e, QuestionImportError):
            logger.error(
                "Import parse error serial=%s (%s): %s",
                serial,
                source_label,
                e,
            )
            summary = traceback.extract_tb(e.__traceback__)
        else:
            logger.error(
                "Unexpected error serial=%s (%s): %s",
                serial,
                source_label,
                e,
            )
            summary = traceback.extract_tb(e.__traceback__)
        if summary:
            logger.error("Error info: %s %s", summary[-1], e)
        return False


def is_question_only_docx(docx_path: str) -> bool:
    try:
        doc = Document(docx_path)
    except Exception:
        return False

    for tbl in doc.tables:
        if not tbl.rows:
            continue

        labels = set()
        for row in tbl.rows:
            if not row.cells:
                continue
            label = " ".join((row.cells[0].text or "").strip().split())
            if label:
                labels.add(label)

        if (
            "Q#:" in labels
            and "Serial #:" in labels
            and "Stem" in labels
            and "Ans:" in labels
        ):
            return True

    return False


def is_explanation_update_format(docx_path: str) -> bool:
    try:
        doc = Document(docx_path)
    except Exception:
        return False

    for _, table in enumerate(doc.tables):
        if len(table.rows) == 5 and len(table.columns) in (2, 3):
            qnum_text = (table.cell(0, 0).text or "").strip()
            if re.match(r"^\D*(\d+)\D*$", qnum_text):
                return True

    return False


def _run_docx_import(
    file_name: str,
    file_data: bytes,
    course: Course,
    create_required: bool,
    auto_verify: bool,
) -> tuple[int, int]:
    success_count = 0
    failure_count = 0
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
        temp_file.write(file_data)
        path = temp_file.name
    try:
        if is_explanation_update_format(path):
            parsed_updates = parse_explanation_updates(path, file_name)
            for question_data in parsed_updates:
                inserted = _insert_question_with_logging(
                    question_data,
                    "docx-explanation-update",
                    lambda qd=question_data: update_question_explanation(qd),
                )
                success_count += int(inserted)
                failure_count += int(not inserted)
        elif is_question_only_docx(path):
            parsed_questions = parse(path)
            for question_data in parsed_questions:
                inserted = _insert_question_with_logging(
                    question_data,
                    "docx-v3",
                    lambda qd=question_data: insert_docx_data_v3(
                        qd, course, create_required
                    ),
                )
                success_count += int(inserted)
                failure_count += int(not inserted)
        else:
            for question_data in parse_questions_from_docx(path, docx_table_format_a):
                inserted = _insert_question_with_logging(
                    question_data,
                    "docx",
                    lambda qd=question_data: insert_docx_data(
                        qd, course, create_required, path, auto_verify
                    ),
                )
                success_count += int(inserted)
                failure_count += int(not inserted)
    finally:
        try:
            os.unlink(path)
        except OSError as e:
            logger.warning("Failed to delete temporary docx file %s: %s", path, e)
    return success_count, failure_count


def _run_csv_import(
    file_name: str,
    file_data: bytes,
    course: Course,
    create_required: bool,
    auto_verify: bool,
) -> tuple[int, int]:
    success_count = 0
    failure_count = 0
    with tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="wb") as temp_file:
        temp_file.write(file_data)
        temp_file.flush()
        path = temp_file.name
    try:
        for question_data in parse_questions_from_csv(path):
            inserted = _insert_question_with_logging(
                question_data,
                "csv",
                lambda qd=question_data: insert_csv_data(
                    qd, course, create_required, auto_verify
                ),
            )
            success_count += int(inserted)
            failure_count += int(not inserted)
    finally:
        try:
            os.unlink(path)
        except OSError as e:
            logger.warning("Failed to delete temporary csv file %s: %s", path, e)
    return success_count, failure_count


_IMPORT_RUNNERS = (
    (".docx", _run_docx_import),
    (".csv", _run_csv_import),
)


def _suffix_for_file_name(file_name: str) -> str | None:
    lower = file_name.lower()
    for suffix, _ in _IMPORT_RUNNERS:
        if lower.endswith(suffix):
            return suffix
    return None


@shared_task
def parse_file(
    file_name: str,
    file_data: bytes,
    course_data: dict,
    uploading_user_id: int,
    create_required: bool,
    upload_result_id: str,
) -> None:
    """
    Celery task to parse uploaded question bank files.

    :param file_name: Name of the uploaded file (extension selects the parser).
    :param file_data: Byte content of the uploaded file.
    :param course_data: Dictionary containing course identifiers (code, year, and semester).
    :param uploading_user_id: ID of the user uploading the file; used for auto-verify.
    :param create_required: Create all required units and subtopics if they don't exist.
    :param upload_result_id: ID of the QuestionUploadResult record to update.
    """
    upload_result = get_upload_result(upload_result_id)
    auto_verify = can_auto_verify(uploading_user_id, course_data)
    try:
        course = Course.objects.get(**course_data)
    except Course.DoesNotExist:
        upload_result.result = upload_result.QuestionUploadResultChoices.FAILURE
        upload_result.progress = 1.0
        upload_result.save(update_fields=["result", "progress"])
        raise ValueError(
            f"Course with code {course_data.get('code')}, year {course_data.get('year')}, "
            f"semester {course_data.get('semester')} does not exist."
        )

    suffix = _suffix_for_file_name(file_name)
    if suffix is None:
        upload_result.result = upload_result.QuestionUploadResultChoices.FAILURE
        upload_result.progress = 1.0
        upload_result.save(update_fields=["result", "progress"])
        raise ValueError(
            "Unsupported file format. Only .docx and .csv files are supported."
        )

    try:
        for registered_suffix, runner in _IMPORT_RUNNERS:
            if suffix == registered_suffix:
                success_count, failure_count = runner(
                    file_name, file_data, course, create_required, auto_verify
                )
                finish_upload_result(upload_result, success_count, failure_count)
                return
    except Exception:
        upload_result.result = upload_result.QuestionUploadResultChoices.FAILURE
        upload_result.progress = 1.0
        upload_result.save(update_fields=["result", "progress"])
        raise


def can_auto_verify(user_id: int, course: dict) -> bool:
    """Check if the user has instructor privileges for the given course."""
    return Enrolment.objects.filter(
        user__id=user_id,
        course__code=course.get("code"),
        course__year=course.get("year"),
        course__semester=course.get("semester"),
        is_instructor=True,
    ).exists()


def parse_select_frequency(value) -> float:
    if value is None:
        return 0.0
    try:
        match = decimal_pattern.search(str(value))
        if match:
            return float(match.group())
    except (ValueError, TypeError):
        pass
    return 0.0


def insert_docx_data(
    question_data: dict,
    course: Course,
    create_required: bool,
    temp_file_name: str,
    auto_verify: bool,
) -> None:
    """
    Inserts parsed question data from a DOCX upload.
    """
    try:
        answer_index = ord(question_data.get("answer").upper().rstrip("() .")) - ord(
            "A"
        )
    except Exception:
        raise QuestionImportError(
            f"Invalid answer format: {question_data.get('answer')}"
        )

    raw_selection_frequencies = question_data.get("option_selection_frequencies", [])
    if len(raw_selection_frequencies) <= answer_index:
        raise QuestionImportError(f"Invalid answer index {answer_index}")

    selection_frequencies = [
        parse_select_frequency(freq) for freq in raw_selection_frequencies
    ]
    selection_frequency = selection_frequencies[answer_index]

    unit_name = question_data.get("unit").strip()
    subtopic_name = question_data.get("subtopic").strip()
    raw_unit_number = question_data.get("unit_number").strip()
    unit_number = int(raw_unit_number) if raw_unit_number not in (None, "") else -1

    with transaction.atomic():
        if create_required:
            unit, _ = Unit.objects.get_or_create(
                defaults={"number": unit_number}, course=course, name=unit_name
            )
            subtopic, _ = UnitSubtopic.objects.get_or_create(
                unit=unit, name=subtopic_name
            )
        else:
            unit = Unit.objects.get(course=course, name=unit_name)
            subtopic = UnitSubtopic.objects.get(unit=unit, name=subtopic_name)

        question = create_question(
            question_data, selection_frequency, subtopic, auto_verify
        )
        created_images = save_images(question_data, question.public_id, temp_file_name)
        question.images.set(created_images)

        create_question_options(question_data, answer_index, question)
        create_question_comments(question_data, question)


def save_v3_images(image_dicts, question_public_id):
    saved_by_digest = {}
    saved_by_ref = {}

    for image in image_dicts:
        image_name = image.get("name", "image.bin")
        image_bytes = image.get("bytes")
        image_ref = image.get("ref")
        if not image_bytes or not image_ref:
            continue

        digest = hashlib.sha256(image_bytes).hexdigest()

        if digest not in saved_by_digest:
            extension = os.path.splitext(image_name)[1].lower() or ".bin"
            filename = f"{question_public_id}_{digest}{extension}"
            saved_by_digest[digest] = QuestionImage.objects.create(
                image_file=ImageFile(BytesIO(image_bytes), name=filename),
                alt_text="",
            )

        saved_by_ref[image_ref] = saved_by_digest[digest]

    return saved_by_ref


def replace_image_placeholders(html, saved_by_ref):
    def repl(match):
        ref = match.group(0)
        img = saved_by_ref.get(ref)

        if not img:
            print("NO IMAGE FOR PLACEHOLDER:", ref)
            return f"<p>[Missing image: {ref}]</p>"

        return f'<img src="/media/{img.image_file.name}">'

    return re.sub(r"\[\[IMG:[^\]]+\]\]", repl, html)


def resolve_question_by_base_serial(base_serial: str) -> Question:
    matches = Question.objects.filter(
        Q(serial_number=base_serial) | Q(serial_number__startswith=base_serial + "_")
    )

    count = matches.count()
    if count == 0:
        raise Question.DoesNotExist(f"No question found for base serial {base_serial}")
    if count > 1:
        raise Question.MultipleObjectsReturned(
            f"Multiple questions found for base serial {base_serial}"
        )

    return matches.get()


def insert_docx_data_v3(
    question_data: dict, course: Course, create_required: bool
) -> None:
    raw_unit_number = question_data.get("unit_number")
    try:
        unit_number = int(raw_unit_number)
    except (TypeError, ValueError):
        unit_number = -1
    unit_name = question_data.get("unit_name") or "Imported Unit"
    subtopic_name = question_data.get("subtopic_name") or "Imported Subtopic"

    with transaction.atomic():
        if create_required:
            unit, _ = Unit.objects.get_or_create(
                course=course,
                number=unit_number,
                defaults={"name": unit_name},
            )
            subtopic, _ = UnitSubtopic.objects.get_or_create(
                unit=unit,
                name=subtopic_name,
            )
        else:
            unit = Unit.objects.get(course=course, name=unit_name)
            subtopic = UnitSubtopic.objects.get(unit=unit, name=subtopic_name)

        question = Question.objects.create(
            subtopic=subtopic,
            serial_number=question_data.get("serial_number")
            or f"{course.code}-DOCX-{question_data['number']}",
            content="",
            answer_explanation="",
            selection_frequency=0,
            difficulty=question_data.get("difficulty", 0),
        )

        content_images_by_ref = save_v3_images(
            question_data.get("content_images", []),
            question.public_id,
        )
        explanation_images_by_ref = save_v3_images(
            question_data.get("explanation_images", []),
            question.public_id,
        )

        question_image_objs = list(content_images_by_ref.values()) + list(
            explanation_images_by_ref.values()
        )

        seen_ids = set()
        unique_question_images = []
        for img in question_image_objs:
            if img.id not in seen_ids:
                seen_ids.add(img.id)
                unique_question_images.append(img)

        question.images.set(unique_question_images)

        content_html = question_data.get("content", "")
        explanation_html = question_data.get("answer_explanation", "")

        content_html = replace_image_placeholders(content_html, content_images_by_ref)
        explanation_html = replace_image_placeholders(
            explanation_html, explanation_images_by_ref
        )

        question.content = content_html
        question.answer_explanation = explanation_html
        question.save(update_fields=["content", "answer_explanation"])

        for opt in question_data.get("options", []):
            option = QuestionOption.objects.create(
                question=question,
                content="",
                is_answer=opt.get("is_answer", False),
                selection_frequency=opt.get("selection_frequency", 0),
            )

            option_images_by_ref = save_v3_images(
                opt.get("images", []),
                question.public_id,
            )

            option.images.set(list(option_images_by_ref.values()))

            option_html = replace_image_placeholders(
                opt.get("content", ""),
                option_images_by_ref,
            )
            option.content = option_html
            option.save(update_fields=["content"])


def update_question_explanation(question_data: dict) -> None:
    question = resolve_question_by_base_serial(question_data["serial_number"])

    explanation_images_by_ref = save_v3_images(
        question_data.get("explanation_images", []),
        question.public_id,
    )

    if explanation_images_by_ref:
        question.images.add(*list(explanation_images_by_ref.values()))

    explanation_html = replace_image_placeholders(
        question_data.get("answer_explanation", ""),
        explanation_images_by_ref,
    )

    question.answer_explanation = explanation_html
    question.save(update_fields=["answer_explanation"])


# TODO: verify with Josh if mismatches should be fixed in the CSV instead
def _resolve_csv_subtopic(
    course: Course,
    unit_tag: str,
    subtopic_tag: str,
    create_required: bool,
    unit_number: int | None,
) -> UnitSubtopic | None:
    """
    Resolve UnitSubtopic from CSV tags; prefer matching subtopic over unit when they disagree.
    If create_required and both tags are set but nothing matches, create unit + subtopic.
    """
    subtopic = None
    if subtopic_tag:
        subtopic_qs = UnitSubtopic.objects.filter(
            unit__course=course, tag=subtopic_tag
        )
        if unit_tag:
            exact = subtopic_qs.filter(unit__tag=unit_tag).first()
            if exact:
                subtopic = exact
            else:
                fallback = subtopic_qs.first()
                if fallback:
                    logger.warning(
                        "subtopic %r not found in unit %r; using unit %r instead",
                        subtopic_tag,
                        unit_tag,
                        fallback.unit.tag,
                    )
                    subtopic = fallback
                elif not create_required:
                    logger.warning(
                        "no subtopic with tag %r for course %s; skipping subtopic",
                        subtopic_tag,
                        course,
                    )
        else:
            subtopic = subtopic_qs.first()
            if not subtopic and not create_required:
                logger.warning(
                    "no subtopic with tag %r for course %s; skipping subtopic",
                    subtopic_tag,
                    course,
                )

    if subtopic is None and create_required and unit_tag and subtopic_tag:
        unum = unit_number if isinstance(unit_number, int) else -1
        unit = Unit.objects.filter(course=course, tag=unit_tag).first()
        if unit is None:
            unit, _ = Unit.objects.get_or_create(
                course=course,
                name=unit_tag,
                defaults={"number": unum, "tag": unit_tag, "description": ""},
            )
        subtopic = UnitSubtopic.objects.filter(unit=unit, tag=subtopic_tag).first()
        if subtopic is None:
            subtopic, _ = UnitSubtopic.objects.get_or_create(
                unit=unit,
                name=subtopic_tag,
                defaults={"tag": subtopic_tag, "description": ""},
            )

    return subtopic


def insert_csv_data(
    question_data: dict,
    course: Course,
    create_required: bool,
    auto_verify: bool,
) -> None:
    """
    Inserts parsed question data from a Brightspace CSV upload.
    """
    answer_letter = (question_data.get("answer") or "").upper()
    if not answer_letter or not ("A" <= answer_letter[0] <= "Z"):
        raise QuestionImportError(f"Invalid answer: {question_data.get('answer')}")
    answer_index = ord(answer_letter[0]) - ord("A")

    options = question_data.get("options") or []
    if answer_index < 0 or answer_index >= len(options):
        raise QuestionImportError(
            f"Answer index {answer_index} out of range for {len(options)} options"
        )

    unit_tag = question_data.get("unit_tag", "").strip()
    subtopic_tag = question_data.get("subtopic_tag", "").strip()
    difficulty = question_data.get("difficulty")
    unit_number = question_data.get("unit_number")
    if unit_number is not None and not isinstance(unit_number, int):
        try:
            unit_number = int(unit_number)
        except (TypeError, ValueError):
            unit_number = None

    with transaction.atomic():
        subtopic = _resolve_csv_subtopic(
            course, unit_tag, subtopic_tag, create_required, unit_number
        )

        question = create_question(
            question_data,
            selection_frequency=0.0,
            subtopic=subtopic,
            is_verified=auto_verify,
            difficulty=difficulty,
        )

        # TODO: image handling (no question images in csv so for now we'll skip implementing it)

        create_question_options(question_data, answer_index, question)
        create_question_comments(question_data, question)


def create_question(
    question_data,
    selection_frequency: float,
    subtopic,
    is_verified: bool,
    difficulty: float = None,
):
    serial_number = (question_data.get("serial_number") or "N/A").strip()
    content = question_data.get("content")
    if not content or str(content).strip() == "":
        raise ValueError(
            f"Question content is empty for question with serial number {serial_number}"
        )

    answer_explanation = question_data.get("explanation")
    if not answer_explanation:
        answer_explanation = ""
    question = Question.objects.create(
        subtopic=subtopic,
        serial_number=question_data.get("serial_number"),
        content=str(content).strip(),
        answer_explanation=str(answer_explanation).strip(),
        selection_frequency=selection_frequency,
        difficulty=difficulty if difficulty is not None else calculate_difficulty_for_test(selection_frequency),
        is_verified=is_verified,
    )
    return question


def create_question_comments(question_data, question):
    comment_text = question_data.get("comments")
    if comment_text is not None and str(comment_text).strip() != "":
        QuestionComment.objects.create(question=question, comment_text=comment_text)


def create_question_options(question_data, answer_index, question):
    """
    Create options for a question. Works for DOCX (with per-option frequencies) and CSV
    (omit or shorten option_selection_frequencies to use model default 0).
    """
    options = question_data.get("options", [])
    freqs = question_data.get("option_selection_frequencies")
    option_explanations = question_data.get("option_explanations") or []
    for idx, option_content in enumerate(options):
        is_answer = idx == answer_index
        if freqs is not None and idx < len(freqs):
            option_selection_frequency = parse_select_frequency(freqs[idx])
        else:
            option_selection_frequency = 0.0
        opt_explanation = ""
        if idx < len(option_explanations) and option_explanations[idx] is not None:
            opt_explanation = str(option_explanations[idx]).strip()
        QuestionOption.objects.create(
            question=question,
            content=option_content,
            explanation=opt_explanation,
            selection_frequency=option_selection_frequency,
            is_answer=is_answer,
        )


def save_images(question_data, question_public_id, file_name):
    question_images = question_data.get("images", [])
    saved_images = []
    for image in question_images:
        image_src = image.get("src")
        image_alt = image.get("alt", "")
        image_ref = image.get("ref")
        with ZipFile(file_name) as docx_zip:
            image_data = docx_zip.read(f"word/{image_src}")
            image_filename = f"{question_public_id}_{image_ref}"
            print(f"Saving image {image_filename}...")
            question_image = QuestionImage.objects.create(
                image_file=ImageFile(BytesIO(image_data), name=image_filename),
                alt_text=image_alt,
            )
            saved_images.append(question_image)
    return saved_images


def calculate_difficulty_for_test(selection_frequency: float) -> float:
    if selection_frequency <= 0 or selection_frequency >= 1:
        return 0.0
    try:
        difficulty = -log((1 / selection_frequency) - 1)
        return round(difficulty, 4)
    except ValueError:
        return 0.0
